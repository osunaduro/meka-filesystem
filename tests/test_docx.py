"""Unit tests for workspace.core.ops.read_docx, write_docx, and edit_docx."""

import pytest
from docx import Document

from workspace.core.errors import AmbiguousMatchError, NoMatchError
from workspace.core.ops import edit_docx, read_docx, write_docx


def test_write_docx_then_read_outline_roundtrips_structure(tmp_path):
    markdown = "# Title\n\nA paragraph with **bold** and *italic* text.\n\n- item one\n- item two\n"

    write_docx(tmp_path, "doc.docx", markdown)
    result = read_docx(tmp_path, "doc.docx")

    styles = [p.style for p in result.paragraphs]
    texts = [p.text for p in result.paragraphs]

    assert "Heading 1" in styles
    assert "Title" in texts
    assert any("bold" in t and "italic" in t for t in texts)
    assert "item one" in texts
    assert "item two" in texts
    assert result.tables == []


def test_write_docx_bullets_use_list_bullet_style(tmp_path):
    write_docx(tmp_path, "doc.docx", "- first\n- second\n")

    result = read_docx(tmp_path, "doc.docx")

    bullet_paragraphs = [p for p in result.paragraphs if p.text in ("first", "second")]
    assert len(bullet_paragraphs) == 2
    assert all(p.style == "List Bullet" for p in bullet_paragraphs)


def test_edit_docx_replaces_matched_text(tmp_path):
    write_docx(tmp_path, "doc.docx", "Hello world\n")

    result = edit_docx(tmp_path, "doc.docx", "Hello", "Goodbye")

    assert result.applied is True
    assert result.occurrences == 1
    outline = read_docx(tmp_path, "doc.docx")
    assert "Goodbye world" in [p.text for p in outline.paragraphs]


def test_edit_docx_dry_run_does_not_write(tmp_path):
    write_docx(tmp_path, "doc.docx", "Hello world\n")

    result = edit_docx(tmp_path, "doc.docx", "Hello", "Goodbye", dry_run=True)

    assert result.applied is False
    outline = read_docx(tmp_path, "doc.docx")
    assert "Hello world" in [p.text for p in outline.paragraphs]


def test_edit_docx_matches_text_split_across_runs(tmp_path):
    # Word commonly splits one visible sentence into multiple <w:r> runs;
    # build such a paragraph by hand to prove the match still works.
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hello ")
    paragraph.add_run("World")
    document.save(tmp_path / "split.docx")

    result = edit_docx(tmp_path, "split.docx", "Hello World", "Goodbye World")

    assert result.occurrences == 1
    outline = read_docx(tmp_path, "split.docx")
    assert "Goodbye World" in [p.text for p in outline.paragraphs]


def test_edit_docx_searches_table_cells(tmp_path):
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "left"
    table.cell(0, 1).text = "right"
    document.save(tmp_path / "table.docx")

    result = edit_docx(tmp_path, "table.docx", "right", "RIGHT")

    assert result.occurrences == 1
    outline = read_docx(tmp_path, "table.docx")
    assert outline.tables[0].rows == [["left", "RIGHT"]]


def test_edit_docx_no_match_attaches_closest_match(tmp_path):
    write_docx(tmp_path, "doc.docx", "The quick brown fox jumps over the lazy dog\n")

    with pytest.raises(NoMatchError) as excinfo:
        edit_docx(tmp_path, "doc.docx", "The quick  brown fox jumps", "x")  # double space typo

    assert excinfo.value.closest_match is not None
    assert "quick" in excinfo.value.closest_match


def test_edit_docx_ambiguous_match_raises(tmp_path):
    write_docx(tmp_path, "doc.docx", "dup\n\ndup\n")

    with pytest.raises(AmbiguousMatchError):
        edit_docx(tmp_path, "doc.docx", "dup", "single")


def test_edit_docx_expected_occurrences_allows_multiple(tmp_path):
    write_docx(tmp_path, "doc.docx", "dup\n\ndup\n")

    result = edit_docx(tmp_path, "doc.docx", "dup", "single", expected_occurrences=2)

    assert result.occurrences == 2
    outline = read_docx(tmp_path, "doc.docx")
    assert [p.text for p in outline.paragraphs] == ["single", "single"]
