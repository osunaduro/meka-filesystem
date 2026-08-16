"""
MEKA Core SDK

Domain:
    Filesystem

Component:
    Operations - write_docx()

Purpose:
    Create or replace a DOCX file rendered from Markdown-formatted text.
"""

"""
MEKA Metadata

Domain:
    filesystem

Component:
    ops.write_docx

Public API:
    write_docx

Dependencies:
    docx
    pathlib
    workspace.core.ops._markdown_lite
    workspace.internal.path

Thread Safe:
    yes

Pure:
    no
"""

# ==========================================================================
# Imports
# ==========================================================================

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from workspace.core.ops._markdown_lite import Blank, BulletItem, Heading, parse_blocks, parse_inline
from workspace.internal.path import resolve as resolve_path


# ==========================================================================
# Internal helpers
# ==========================================================================

def _add_runs(paragraph: Paragraph, text: str) -> None:
    for run in parse_inline(text):
        added = paragraph.add_run(run.text)
        added.bold = run.bold or None
        added.italic = run.italic or None


# ==========================================================================
# Public API
# ==========================================================================

def write_docx(root: Path, path: str | Path, markdown: str) -> None:
    """Create or replace a DOCX file rendered from Markdown-formatted text.

    Same lightweight subset supported by ``write_pdf``: headings (#, ##,
    ###), bullet lists (- or *), bold (**text**), italic (*text*), and
    plain paragraphs. Not supported: tables, links, images, code blocks,
    numbered lists.

    Parameters
    ----------
    root : Path
        Workspace root.

    path : str | Path
        Relative path inside the workspace for the .docx file to create.

    markdown : str
        Markdown-formatted source text.
    """
    target = resolve_path(root, path)
    document = Document()

    for block in parse_blocks(markdown):
        if isinstance(block, Blank):
            continue
        elif isinstance(block, Heading):
            _add_runs(document.add_heading("", level=block.level), block.text)
        elif isinstance(block, BulletItem):
            _add_runs(document.add_paragraph("", style="List Bullet"), block.text)
        else:
            _add_runs(document.add_paragraph(""), block.text)

    document.save(target)
