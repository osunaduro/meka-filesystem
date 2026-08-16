"""
MEKA Core SDK

Domain:
    Filesystem

Component:
    Operations - edit_docx()

Purpose:
    Replace text matched by content within a DOCX file's paragraphs and
    table cells.
"""

"""
MEKA Metadata

Domain:
    filesystem

Component:
    ops.edit_docx

Public API:
    edit_docx

Dependencies:
    difflib
    docx
    pathlib
    workspace.core.errors
    workspace.core.models
    workspace.core.ops._fuzzy_match
    workspace.core.ops.read_docx
    workspace.internal.path

Thread Safe:
    yes

Pure:
    no
"""

# ==========================================================================
# Imports
# ==========================================================================

import difflib
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from workspace.core.errors import AmbiguousMatchError, NoMatchError
from workspace.core.models import EditResult
from workspace.core.ops._docx_walk import iter_block_items
from workspace.core.ops._fuzzy_match import closest_match
from workspace.internal.path import resolve as resolve_path


# ==========================================================================
# Internal helpers
# ==========================================================================

def _editable_paragraphs(document: Document) -> list[Paragraph]:
    """Every paragraph in the document, including those inside table cells.

    Only cells' own paragraphs are visited (not nested tables), matching
    the shapes ``write_docx`` produces.
    """
    paragraphs: list[Paragraph] = []
    for item in iter_block_items(document):
        if isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        else:
            paragraphs.append(item)
    return paragraphs


def _rewrite_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace a paragraph's content with a single run of ``text``.

    Simplification versus editing raw XML: Word commonly splits a
    paragraph's visible text across several runs (spell-check, formatting
    boundaries), so a content match can cross run boundaries. Rather than
    mapping the replacement back onto the original runs, the whole
    paragraph is rewritten as one run — this can flatten run-level
    formatting (e.g. a single bolded word) within that specific
    paragraph, which is a deliberate, documented trade-off.
    """
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


# ==========================================================================
# Public API
# ==========================================================================

def edit_docx(
    root: Path,
    path: str | Path,
    old_text: str,
    new_text: str,
    *,
    expected_occurrences: int = 1,
    dry_run: bool = False,
) -> EditResult:
    """Replace text matched by content within a DOCX file.

    Searches the text of every paragraph and table cell (each paragraph's
    text is matched independently — ``old_text`` must fall entirely
    within a single paragraph, it cannot span two).

    Parameters
    ----------
    root : Path
        Workspace root.

    path : str | Path
        Relative path inside the workspace to a .docx file.

    old_text : str
        Exact text to search for within a paragraph's content.

    new_text : str
        Replacement text.

    expected_occurrences : int, optional
        Number of times ``old_text`` must appear across the document for
        the edit to proceed. Defaults to 1.

    dry_run : bool, optional
        When True, compute and return the diff without writing to disk.

    Returns
    -------
    EditResult
        ``diff`` is a unified diff over the document's paragraph text,
        one paragraph per line.

    Raises
    ------
    NoMatchError
        ``old_text`` was not found in any paragraph. Carries a fuzzy
        ``closest_match``/``similarity`` hint when a near-miss exists.

    AmbiguousMatchError
        ``old_text`` matched a different number of times than
        ``expected_occurrences``.
    """
    if expected_occurrences < 1:
        raise ValueError("'expected_occurrences' must be greater than zero.")

    target = resolve_path(root, path)
    document = Document(target)
    paragraphs = _editable_paragraphs(document)

    original_lines = [paragraph.text for paragraph in paragraphs]
    occurrences = sum(line.count(old_text) for line in original_lines)

    if occurrences == 0:
        window, ratio = closest_match("\n".join(original_lines), old_text)
        raise NoMatchError(
            f"No match for the requested text in {path}.",
            closest_match=window,
            similarity=ratio,
        )

    if occurrences != expected_occurrences:
        raise AmbiguousMatchError(
            f"Expected {expected_occurrences} occurrence(s) of the requested "
            f"text in {path}, found {occurrences}."
        )

    updated_lines = list(original_lines)
    for index, paragraph in enumerate(paragraphs):
        if old_text in original_lines[index]:
            updated_lines[index] = original_lines[index].replace(old_text, new_text)
            if not dry_run:
                _rewrite_paragraph_text(paragraph, updated_lines[index])

    diff = "".join(
        difflib.unified_diff(
            [line + "\n" for line in original_lines],
            [line + "\n" for line in updated_lines],
            fromfile=str(path),
            tofile=str(path),
        )
    )

    if not dry_run:
        document.save(target)

    return EditResult(path=target, applied=not dry_run, occurrences=occurrences, diff=diff)
